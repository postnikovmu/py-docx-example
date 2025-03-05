import docx
from lxml import etree
import zipfile


def get_document_xml(docx_filename):
    with zipfile.ZipFile(docx_filename) as zip_ref:
        xml_content = zip_ref.read('word/document.xml')
    return xml_content


def get_theme_xml(docx_filename):
    with zipfile.ZipFile(docx_filename) as zip_ref:
        xml_content = zip_ref.read('word/theme/theme1.xml')
    return xml_content


def get_styles_xml(docx_filename):
    with zipfile.ZipFile(docx_filename) as zip_ref:
        xml_content = zip_ref.read('word/styles.xml')
    return xml_content


def parse_xml(xml_string):
    return etree.fromstring(xml_string)


def print_text_elements(root):
    for elem in root.iter():
        if elem.tag.endswith('t'):
            if elem.text is not None:
                print("Текст из XML:", elem.text)


def get_theme_fonts(root):
    major_font = None
    minor_font = None
    for elem in root.iter():
        if elem.tag.endswith('majorFont'):
            for child in elem:
                if child.tag.endswith('latin'):
                    major_font = child.attrib['typeface']
        elif elem.tag.endswith('minorFont'):
            for child in elem:
                if child.tag.endswith('latin'):
                    minor_font = child.attrib['typeface']

    return major_font, minor_font


def get_default_font_size(root):
    default_font_size = None
    ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
    for elem in root.iter():
        if elem.tag.endswith('rPrDefault'):
            for child_btw in elem:
                if child_btw.tag.endswith('rPr'):
                    for child in child_btw:
                        if child.tag == ns + 'sz':
                            default_font_size = str(float(child.attrib.get(f'{ns}val'))/2)
                            print(f"Значение {type(default_font_size)} для w:sz: {default_font_size}")

    return default_font_size


def get_file_content_python_docx(file_path: str) -> None:
    """
    This function based on the python-docx library.
    It can print paragraphs one by one.
    The paragraph can contain more then one 'Run' object,
    in case there are several different formattings inside the paragraph.

    ! limitation
    Can't get formatting for old *.doc files
    The primary reason python-docx cannot retrieve
    formatting from old .doc files is due to the inherent differences between the two formats.
    The .doc format uses a binary structure that encapsulates text and formatting in a way that is not compatible
    with the XML-based approach of .docx. This means that any code written for python-docx,
    which relies on parsing XML elements, will not function correctly when applied to .doc files

    :param file_path:
    :return:
    """
    try:
        # Open the document
        doc = docx.Document(file_path)

        default_style = doc.styles['Normal']
        print(dir(default_style))
        print('Имя шрифта по умолчанию:', default_style.font.name)
        print('Размер шрифта по умолчанию:', default_style.font.size)

        for style in doc.styles:
            print(style.name, style.style_id, style.type, style.font.name if hasattr(style, 'font') else None)

            #font = style.font

            # Получение параметров шрифта
            #font_name = font.name  # Название шрифта
            #font_size = font.size  # Размер шрифта (в пунктах)
            #font_bold = font.bold  # Жирный текст (True/False)
            #font_italic = font.italic  # Курсив (True/False)
            #font_underline = font.underline  # Подчеркивание (None/Single/Double)
            #print(font_name, font_size, font_bold, font_italic, font_underline)

        # Iterate through all elements and print their text
        index_for_paragraphs = None
        index_for_tables = None

        for element_num in range(len(doc.element.body.inner_content_elements)):

            element = doc.element.body.inner_content_elements[element_num]
            element_type = None
            if element.tag.endswith('p'):  #check for a paragraph
                element_type = 'paragraph'
                if index_for_paragraphs is None:
                    index_for_paragraphs = 0
                else:
                    index_for_paragraphs += 1

            elif element.tag.endswith('tbl'):  #check for a paragraph
                element_type = 'table'
                if index_for_tables is None:
                    index_for_tables = 0
                else:
                    index_for_tables += 1

            if element_type == 'paragraph':
                element_content = doc.paragraphs[index_for_paragraphs]
                print(element_content.text)

                for run in element_content.runs:
                    text = run.text
                    is_bold = run.bold if run.bold is not None else False  # Default to False if None
                    is_italic = run.italic if run.italic is not None else False  # Default to False if None
                    font_size = str(run.font.size.pt) if run.font.size else 'Not Set'
                    font_color = str(run.font.color.rgb) if run.font.color and run.font.color.rgb else 'Not Set'
                    font_name = str(run.font.name) if run.font.name else 'Not Set'

                    print(f'Run Text: {text}')
                    print(f'Run Bold: {is_bold}')
                    print(f'Run Italic: {is_italic}')
                    print(f'Run Font Size: {font_size}')
                    print(f'Run Font Color: {font_color}')
                    print(f'Run Font Name: {font_name}\n')

            if element_type == 'table':
                print('>>>>>>Table')
                element_content = doc.tables[index_for_tables]
                for row in element_content.rows:
                    for cell in row.cells:
                        # print data from the cell
                        for paragraph in cell.paragraphs:
                            print(paragraph.text)
                print('<<<<<<Table')

        document_xml_content = get_document_xml(file_path)
        theme_xml_content = get_theme_xml(file_path)
        styles_xml_content = get_styles_xml(file_path)

        document_parsed = parse_xml(document_xml_content)
        theme_parsed = parse_xml(theme_xml_content)
        styles_parsed = parse_xml(styles_xml_content)

        print_text_elements(document_parsed)
        major_font, minor_font = get_theme_fonts(theme_parsed)
        default_font_size = get_default_font_size(styles_parsed)
        print(f"Major Font: {major_font}")
        print(f"Minor Font: {minor_font}")
        print(f"Default Font Size: {default_font_size}")

    except IOError:
        print('There was an error opening the file!')

